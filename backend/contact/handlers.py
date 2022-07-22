from docx import Document


def get_donation_data(docx_file):
    data = {}
    document = Document(docx_file)
    data.update({'title': document.paragraphs[0].text})
    for row in document.tables[0].rows:
        header, body = row.cells
        data.update({header.text: body.text})
    return data
